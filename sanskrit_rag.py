# Sanskrit Medical Texts RAG System - Complete Implementation
# File: sanskrit_rag_system.py

import os
import torch
import warnings
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
import logging
from pathlib import Path
import json
import re

# Core libraries
from transformers import (
    AutoTokenizer, 
    AutoModelForCausalLM, 
    BitsAndBytesConfig,
    pipeline
)

from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import PyPDF2
import docx
import gradio as gr

# Suppress warnings
warnings.filterwarnings("ignore")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class RAGConfig:
    """Configuration for the RAG system"""
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"
    llm_model: str = "mistralai/Mistral-7B-Instruct-v0.2"
    chunk_size: int = 800
    chunk_overlap: int = 100
    top_k_retrieval: int = 4
    max_new_tokens: int = 512
    temperature: float = 0.7
    db_path: str = "./chroma_db"

class DocumentProcessor:
    """Handle document loading and preprocessing"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100,
            separators=["\n\n", "\n", "।", ".", " "]
        )
    
    def load_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            return ""
    
    def load_txt(self, file_path: str) -> str:
        """Load text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            return ""
    
    def clean_sanskrit_text(self, text: str) -> str:
        """Clean and preprocess Sanskrit text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep Devanagari
        text = re.sub(r'[^\w\s\u0900-\u097F।\.]', ' ', text)
        # Remove empty lines
        text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
        return text.strip()
    
    def process_document(self, file_path: str, doc_name: str = None) -> List[Document]:
        """Process document and return chunks"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Load document based on extension
        if file_path.suffix.lower() == '.pdf':
            text = self.load_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            text = self.load_docx(str(file_path))
        elif file_path.suffix.lower() == '.txt':
            text = self.load_txt(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        if not text.strip():
            raise ValueError(f"No text extracted from {file_path}")
        
        # Clean text
        text = self.clean_sanskrit_text(text)
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create Document objects
        documents = []
        doc_name = doc_name or file_path.stem
        
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only add non-empty chunks
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": str(file_path),
                        "document_name": doc_name,
                        "chunk_id": i,
                        "chunk_size": len(chunk)
                    }
                )
                documents.append(doc)
        
        logger.info(f"Created {len(documents)} chunks from {doc_name}")
        return documents

class VectorStore:
    """Handle vector database operations"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.embedding_model = SentenceTransformer(config.embedding_model)
        self.client = chromadb.PersistentClient(path=config.db_path)
        self.collection = self.client.get_or_create_collection(
            name="sanskrit_medical_texts",
            metadata={"description": "Sanskrit medical texts like Charak Samhita"}
        )
        logger.info(f"Vector store initialized with {self.collection.count()} documents")
    
    def add_documents(self, documents: List[Document]):
        """Add documents to vector store"""
        if not documents:
            logger.warning("No documents to add")
            return
        
        texts = [doc.page_content for doc in documents]
        metadatas = [doc.metadata for doc in documents]
        
        # Generate embeddings
        embeddings = self.embedding_model.encode(texts, show_progress_bar=True)
        
        # Create unique IDs
        ids = [f"{doc.metadata['document_name']}_{doc.metadata['chunk_id']}" 
               for doc in documents]
        
        # Add to collection
        self.collection.add(
            embeddings=embeddings.tolist(),
            documents=texts,
            metadatas=metadatas,
            ids=ids
        )
        
        logger.info(f"Added {len(documents)} documents to vector store")
    
    def search(self, query: str, top_k: int = None) -> List[Dict]:
        """Search for relevant documents"""
        top_k = top_k or self.config.top_k_retrieval
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query])
        
        # Search in collection
        results = self.collection.query(
            query_embeddings=query_embedding.tolist(),
            n_results=top_k
        )
        
        # Format results
        search_results = []
        for i in range(len(results['documents'][0])):
            search_results.append({
                'content': results['documents'][0][i],
                'metadata': results['metadatas'][0][i],
                'score': 1 - results['distances'][0][i]  # Convert distance to similarity
            })
        
        return search_results
    
    def get_stats(self) -> Dict:
        """Get database statistics"""
        return {
            "total_documents": self.collection.count(),
            "embedding_model": self.config.embedding_model
        }

class LLMManager:
    """Handle LLM operations"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.tokenizer = None
        self.model = None
        self.pipeline = None
        self._setup_model()
    
    def _setup_model(self):
        """Initialize the LLM with quantization"""
        logger.info(f"Loading model: {self.config.llm_model}")
        
        # Configure quantization for memory efficiency
        bnb_config = BitsAndBytesConfig(
            load_in_4bit=True,
            bnb_4bit_quant_type="nf4",
            bnb_4bit_compute_dtype=torch.float16,
            bnb_4bit_use_double_quant=True
        )
        
        try:
            # Load tokenizer
            self.tokenizer = AutoTokenizer.from_pretrained(
                self.config.llm_model,
                trust_remote_code=True
            )
            if self.tokenizer.pad_token is None:
                self.tokenizer.pad_token = self.tokenizer.eos_token
            
            # Load model
            self.model = AutoModelForCausalLM.from_pretrained(
                self.config.llm_model,
                quantization_config=bnb_config,
                device_map="auto",
                trust_remote_code=True
            )
            
            # Create pipeline
            self.pipeline = pipeline(
                "text-generation",
                model=self.model,
                tokenizer=self.tokenizer,
                max_new_tokens=self.config.max_new_tokens,
                temperature=self.config.temperature,
                do_sample=True,
                pad_token_id=self.tokenizer.eos_token_id
            )
            
            logger.info("Model loaded successfully")
            
        except Exception as e:
            logger.error(f"Error loading model: {e}")
            raise
    
    def generate_response(self, prompt: str) -> str:
        """Generate response using the LLM"""
        try:
            response = self.pipeline(prompt)
            generated_text = response[0]['generated_text']
            
            # Extract only the new generated part
            if prompt in generated_text:
                answer = generated_text[len(prompt):].strip()
            else:
                answer = generated_text.strip()
            
            return answer
            
        except Exception as e:
            logger.error(f"Error generating response: {e}")
            return "I apologize, but I encountered an error while generating the response."

class SanskritRAGAgent:
    """Main RAG agent for Sanskrit medical texts"""
    
    def __init__(self, config: RAGConfig = None):
        self.config = config or RAGConfig()
        self.doc_processor = DocumentProcessor()
        self.vector_store = VectorStore(self.config)
        self.llm_manager = LLMManager(self.config)
        self.conversation_history = []
        
        logger.info("Sanskrit RAG Agent initialized successfully")
    
    def add_document(self, file_path: str, doc_name: str = None) -> Dict:
        """Add a new document to the knowledge base"""
        try:
            documents = self.doc_processor.process_document(file_path, doc_name)
            self.vector_store.add_documents(documents)
            
            return {
                "status": "success",
                "message": f"Successfully added {len(documents)} chunks from {doc_name or Path(file_path).stem}",
                "chunks_added": len(documents)
            }
        
        except Exception as e:
            logger.error(f"Error adding document: {e}")
            return {
                "status": "error",
                "message": f"Failed to add document: {str(e)}"
            }
    
    def create_prompt(self, query: str, context_docs: List[Dict]) -> str:
        """Create a well-structured prompt for the LLM"""
        
        context = "\n\n".join([
            f"Reference {i+1} (from {doc['metadata']['document_name']}):\n{doc['content']}"
            for i, doc in enumerate(context_docs)
        ])
        
        prompt = f"""<s>[INST] You are an expert in Ayurvedic medicine and Sanskrit medical texts. You have access to authentic texts like Charak Samhita and other classical works.

Context from Sanskrit Medical Texts:
{context}

Based on the above context from authentic Sanskrit medical texts, please answer the following question about Ayurvedic medicine:

Question: {query}

Instructions:
1. Provide accurate information based on the given context
2. If the context contains relevant Sanskrit terms or concepts, explain them
3. Reference the source text when possible
4. If the context doesn't fully answer the question, state that clearly
5. Maintain the traditional Ayurvedic perspective while making it understandable

Answer: [/INST]"""
        
        return prompt
    
    def query(self, question: str) -> Dict:
        """Process a query and return response with sources"""
        try:
            # Retrieve relevant documents
            search_results = self.vector_store.search(question, self.config.top_k_retrieval)
            
            if not search_results:
                return {
                    "answer": "I couldn't find relevant information in the knowledge base to answer your question.",
                    "sources": [],
                    "confidence": 0.0
                }
            
            # Create prompt with context
            prompt = self.create_prompt(question, search_results)
            
            # Generate response
            answer = self.llm_manager.generate_response(prompt)
            
            # Prepare sources
            sources = [
                {
                    "document": result['metadata']['document_name'],
                    "content_preview": result['content'][:200] + "...",
                    "relevance_score": result['score']
                }
                for result in search_results
            ]
            
            # Calculate confidence based on retrieval scores
            avg_score = sum(r['score'] for r in search_results) / len(search_results)
            confidence = min(avg_score * 1.2, 1.0)  # Scale and cap at 1.0
            
            # Store in conversation history
            self.conversation_history.append({
                "question": question,
                "answer": answer,
                "sources_count": len(sources)
            })
            
            return {
                "answer": answer,
                "sources": sources,
                "confidence": confidence,
                "retrieval_count": len(search_results)
            }
            
        except Exception as e:
            logger.error(f"Error processing query: {e}")
            return {
                "answer": "I apologize, but I encountered an error while processing your question.",
                "sources": [],
                "confidence": 0.0
            }
    
    def get_stats(self) -> Dict:
        """Get system statistics"""
        vector_stats = self.vector_store.get_stats()
        return {
            **vector_stats,
            "conversations": len(self.conversation_history),
            "model": self.config.llm_model,
            "embedding_model": self.config.embedding_model
        }

def create_gradio_interface(agent: SanskritRAGAgent) -> gr.Interface:
    """Create Gradio web interface"""
    
    def query_interface(question: str) -> Tuple[str, str, str]:
        if not question.strip():
            return "Please enter a question.", "", ""
        
        result = agent.query(question)
        
        answer = result['answer']
        confidence = f"Confidence: {result['confidence']:.2%}"
        
        sources_text = "\n\n".join([
            f"Source {i+1}: {source['document']}\n"
            f"Relevance: {source['relevance_score']:.2%}\n"
            f"Preview: {source['content_preview']}"
            for i, source in enumerate(result['sources'])
        ])
        
        return answer, confidence, sources_text
    
    def upload_document(file) -> str:
        if file is None:
            return "No file uploaded."
        
        try:
            result = agent.add_document(file.name)
            return result['message']
        except Exception as e:
            return f"Error uploading document: {str(e)}"
    
    def get_system_stats() -> str:
        stats = agent.get_stats()
        return f"""
        System Statistics:
        - Total Documents: {stats['total_documents']}
        - Conversations: {stats['conversations']}
        - Model: {stats['model']}
        - Embedding Model: {stats['embedding_model']}
        """
    
    # Create interface
    with gr.Blocks(title="Sanskrit Medical Texts RAG System") as demo:
        gr.Markdown("# 🕉️ Sanskrit Medical Texts RAG System\nQuery Charak Samhita and other Ayurvedic texts")
        
        with gr.Tab("Ask Questions"):
            question_input = gr.Textbox(
                label="Your Question",
                placeholder="Ask about Ayurvedic medicine, treatments, or concepts...",
                lines=2
            )
            submit_btn = gr.Button("Ask", variant="primary")
            
            answer_output = gr.Textbox(
                label="Answer",
                lines=10,
                interactive=False
            )
            confidence_output = gr.Textbox(
                label="Confidence",
                lines=1,
                interactive=False
            )
            sources_output = gr.Textbox(
                label="Sources",
                lines=8,
                interactive=False
            )
            
            submit_btn.click(
                query_interface,
                inputs=[question_input],
                outputs=[answer_output, confidence_output, sources_output]
            )
        
        with gr.Tab("Upload Documents"):
            file_input = gr.File(
                label="Upload Document (PDF, DOCX, TXT)",
                file_types=[".pdf", ".docx", ".txt"]
            )
            upload_btn = gr.Button("Upload", variant="primary")
            upload_output = gr.Textbox(
                label="Upload Status",
                lines=3,
                interactive=False
            )
            
            upload_btn.click(
                upload_document,
                inputs=[file_input],
                outputs=[upload_output]
            )
        
        with gr.Tab("System Stats"):
            stats_btn = gr.Button("Refresh Stats", variant="secondary")
            stats_output = gr.Textbox(
                label="System Statistics",
                lines=8,
                interactive=False
            )
            
            stats_btn.click(
                get_system_stats,
                outputs=[stats_output]
            )
    
    return demo

# Main execution
def main():
    """Main function to run the Sanskrit RAG system"""
    
    # Initialize configuration
    config = RAGConfig()
    
    # Create RAG agent
    print("🚀 Initializing Sanskrit RAG System...")
    agent = SanskritRAGAgent(config)
    
    # Example: Add Charak Samhita if file exists
    charak_file = "charak_samhita.pdf"  # Replace with your file path
    if os.path.exists(charak_file):
        print(f"📚 Adding {charak_file} to knowledge base...")
        result = agent.add_document(charak_file, "Charak Samhita")
        print(f"✅ {result['message']}")
    else:
        print(f"⚠️  {charak_file} not found. Please add your documents using the web interface.")
    
    # Create and launch web interface
    print("🌐 Launching web interface...")
    demo = create_gradio_interface(agent)
    demo.launch(
        server_name="0.0.0.0",
        server_port=7860,
        share=False  # Set to True if you want a public link
    )

if __name__ == "__main__":
    main()